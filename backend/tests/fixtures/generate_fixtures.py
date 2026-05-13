"""
Generate test DOCX fixtures for docx-jast integration tests.

Produces 4 articles covering distinct extraction scenarios:
  1. article_full_research.docx   — happy path: title, abstract, 6 sections, 5 IEEE refs
  2. article_bibliography_alias.docx — "Bibliography" heading alias, 4 APA refs
  3. article_no_references.docx   — no references section → parser warning
  4. article_review_paper.docx    — "Works Cited" heading alias, 8 refs, 5 sections
"""

import os
from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH

OUT = os.path.dirname(os.path.abspath(__file__))


# ── helpers ───────────────────────────────────────────────────────────────────

def new_doc():
    doc = Document()
    for sec in doc.sections:
        sec.page_height = Cm(29.7)
        sec.page_width  = Cm(21.0)
        for attr in ("left_margin", "right_margin", "top_margin", "bottom_margin"):
            setattr(sec, attr, Cm(2.54))
    return doc


def p(doc, text, style="Normal", bold=False, italic=False,
      align=WD_ALIGN_PARAGRAPH.LEFT, size=12):
    para = doc.add_paragraph(style=style)
    para.alignment = align
    run = para.add_run(text)
    run.bold = bold
    run.italic = italic
    run.font.name = "Times New Roman"
    run.font.size = Pt(size)
    para.paragraph_format.line_spacing = Pt(size * 1.5)
    return para


def save(doc, name):
    path = os.path.join(OUT, name)
    doc.save(path)
    print(f"  saved: {path}")


# ── Article 1: Full research paper ────────────────────────────────────────────
# Parser expectations: title extracted, abstract (Strategy 2), 5 IEEE refs,
# section_count > 0, zero warnings about title/abstract/refs.

def make_full_research():
    doc = new_doc()
    title = (
        "Quantitative Analysis of Snazzle-Induced Blorfification "
        "in Distributed Quonk Networks: A Zibbit-Based Approach"
    )
    pp = doc.add_paragraph(style="Title")
    pp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = pp.add_run(title)
    r.font.name = "Times New Roman"; r.font.size = Pt(16); r.bold = True

    p(doc, "Iqbal A.¹, Frumpleston B.², Zibbitova C.¹",
      align=WD_ALIGN_PARAGRAPH.CENTER, italic=True)
    p(doc, "¹Institute of Nonsensical Sciences, Blorfville  "
           "²Department of Quonk Engineering, Snazzleton University",
      align=WD_ALIGN_PARAGRAPH.CENTER, size=10)
    doc.add_paragraph()

    # Abstract via Heading1 keyword (parser Strategy 2)
    p(doc, "Abstract", style="Heading1", bold=True)
    p(doc,
      "This paper presents a comprehensive investigation into the blorfification "
      "behaviour exhibited by snazzle-class nodes within heterogeneous quonk "
      "networks. Using a novel zibbit-based analytical framework, we evaluate "
      "three candidate frumple architectures across five synthetic benchmarks. "
      "Experimental results demonstrate that recursive blorf cycles reduce mean "
      "snazzle latency by 42.7% while maintaining quonk coherence above 0.83. "
      "Our findings suggest zibbit-aware scheduling yields statistically "
      "significant gains (p < 0.001) over baseline wumble-first approaches.")
    p(doc, "Keywords: blorf, snazzle, zibbit, quonk, frumple", italic=True, size=11)
    doc.add_paragraph()

    sections = [
        ("1. Introduction", [
            "The proliferation of snazzle-enabled quonk networks has created "
            "unprecedented demand for robust blorfification protocols. Prior work "
            "by Frumpleston et al. [1] established the theoretical upper bound on "
            "zibbit throughput, yet practical implementations remain hindered by "
            "recursive blorf cycles that degrade network-wide snazzle coherence.",
            "Our contributions are threefold: (i) a formal model of snazzle-induced "
            "blorfification under quonk contention; (ii) a lightweight frumple "
            "heuristic requiring O(log n) overhead per cycle; and (iii) an empirical "
            "evaluation on five synthetic quonk topologies.",
        ]),
        ("2. Related Work", [
            "Early work on blorfification focused on static quonk partitioning [2]. "
            "Zibbitova and colleagues demonstrated dynamic frumple rebalancing "
            "reduces blorf-induced jitter by up to 38% [3].",
            "Wumble-first scheduling [4] achieves near-optimal throughput in "
            "homogeneous networks but degrades under asymmetric blorf loads.",
        ]),
        ("3. Methodology", [
            "We model the quonk network as a directed graph G = (V, E) where each "
            "node represents a snazzle-capable blorf unit. Blorfification events "
            "are modelled as Poisson arrivals with rate λ from [1].",
            "Five synthetic benchmarks: (B1) uniform blorf load, (B2) bursty snazzle "
            "traffic, (B3) adversarial quonk partitioning, (B4) mixed frumple sizes, "
            "and (B5) high-churn zibbit topology. Each run 100 times.",
        ]),
        ("4. Results", [
            "The proposed ZFS achieves 42.7% reduction in mean snazzle latency "
            "relative to wumble-first baseline, with gains from 31.2% (B1) to 58.9% (B3).",
            "Quonk coherence remains above 0.87 in all conditions. "
            "Statistical significance confirmed via paired t-tests (p < 0.001).",
        ]),
        ("5. Discussion", [
            "The superiority of ZFS over wumble-first scheduling is attributable to "
            "the real-time frumple feedback loop, which prevents blorf queue buildup.",
        ]),
        ("6. Conclusion", [
            "ZFS reduces mean latency by 42.7%, maintains quonk coherence above "
            "threshold across all benchmarks, and incurs negligible overhead.",
        ]),
    ]
    for heading, paras in sections:
        p(doc, heading, style="Heading1", bold=True)
        for body in paras:
            p(doc, body)
        doc.add_paragraph()

    p(doc, "References", style="Heading1", bold=True)
    for ref in [
        '[1] B. Frumpleston, C. Zibbitova, and A. Iqbal, "Upper bounds on zibbit '
        'throughput," J. Nonsensical Comput., vol. 12, no. 3, pp. 101-118, 2022.',
        '[2] Q. Blorfman and S. Snazzleton, "Static quonk partitioning," in Proc. '
        'ICFE, 2019, pp. 45-52.',
        '[3] C. Zibbitova, M. Wumble, and P. Quonkfield, "Dynamic frumple '
        'rebalancing," IEEE Trans. Quonk Syst., vol. 8, no. 1, pp. 23-37, 2021.',
        '[4] R. Wumble, "Wumble-first scheduling: theory and practice," ACM Comput. '
        'Surv., vol. 54, no. 6, Art. 128, 2023.',
        '[5] A. Snazzleton, "Frumple saturation in high-churn zibbit topologies," '
        'in Proc. IEEE SBN, 2024, pp. 200-207.',
    ]:
        p(doc, ref)

    save(doc, "article_full_research.docx")


# ── Article 2: Bibliography alias ─────────────────────────────────────────────
# Parser expectations: "Bibliography" heading → ref_count=4, has_abstract=true.

def make_bibliography_alias():
    doc = new_doc()
    title = (
        "Digital Strategy Adoption in SME Quonk Enterprises: "
        "A Frumple-Theoretic Perspective"
    )
    pp = doc.add_paragraph(style="Title")
    pp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = pp.add_run(title)
    r.font.name = "Times New Roman"; r.font.size = Pt(16); r.bold = True

    p(doc, "Blorfstein, R., & Snazzleton, M.", align=WD_ALIGN_PARAGRAPH.CENTER, italic=True)
    doc.add_paragraph()

    p(doc, "Abstract", style="Heading1", bold=True)
    p(doc,
      "Small and medium enterprises (SMEs) operating in the quonk sector face "
      "increasing pressure to adopt frumple-theoretic digital strategies. This "
      "paper applies the Zibbit Adoption Model (ZAM) to a sample of 312 SMEs "
      "across three blorf-intensive industries. Structural equation modelling "
      "reveals that snazzle readiness and executive frumple commitment are the "
      "strongest predictors of successful digital transformation. Implications "
      "for quonk policy and managerial practice are discussed.")
    p(doc, "Keywords: digital strategy, frumple theory, SME, quonk, ZAM", italic=True, size=11)
    doc.add_paragraph()

    sections = [
        ("1. Introduction", [
            "Digital transformation has become a strategic imperative for quonk-sector "
            "SMEs. Yet adoption of frumple-theoretic tools remains uneven, with smaller "
            "blorf enterprises lagging behind their larger counterparts (Snazzleton & "
            "Wumble, 2021). This study addresses the gap by applying ZAM to a "
            "cross-industry sample.",
        ]),
        ("2. Theoretical Framework", [
            "The Zibbit Adoption Model extends the Technology Acceptance Model by "
            "incorporating quonk-specific constructs: snazzle readiness, frumple "
            "commitment, and blorf integration capacity. We hypothesise that these "
            "three constructs jointly explain digital strategy adoption variance.",
        ]),
        ("3. Methodology", [
            "Data were collected via structured questionnaire from 312 CEOs of quonk-"
            "sector SMEs across Blorfville, Snazzton, and Zibbitford. Construct "
            "reliability was assessed using Cronbach's alpha (all > 0.80). "
            "Structural equation modelling was performed in R using the lavaan package.",
        ]),
        ("4. Findings and Conclusion", [
            "Snazzle readiness (β = 0.41, p < 0.001) and frumple commitment "
            "(β = 0.38, p < 0.001) are the strongest predictors of adoption. Blorf "
            "integration capacity showed a significant moderating effect. These "
            "findings extend ZAM to the quonk SME context and inform policy design.",
        ]),
    ]
    for heading, paras in sections:
        p(doc, heading, style="Heading1", bold=True)
        for body in paras:
            p(doc, body)
        doc.add_paragraph()

    # "Bibliography" — tests the alias in references.rs
    p(doc, "Bibliography", style="Heading1", bold=True)
    for ref in [
        "Blorfstein, R., & Wumble, T. (2020). Frumple theory and SME performance. "
        "Journal of Quonk Management, 15(2), 88-104.",
        "Snazzleton, M., & Wumble, P. (2021). Digital transformation in blorf-intensive "
        "industries. Business Analytics Review, 9(1), 12-30.",
        "Zibbitova, C. (2019). The Zibbit Adoption Model: A primer. "
        "Academy of Frumple Science, 7(4), 210-225.",
        "Quonkfield, D., & Blorfman, A. (2022). Executive commitment and tech adoption "
        "in SMEs. Strategic Mgmt. J., 43(6), 1150-1172.",
    ]:
        p(doc, ref)

    save(doc, "article_bibliography_alias.docx")


# ── Article 3: No references section ─────────────────────────────────────────
# Parser expectations: reference_count=0, warnings include ref warning,
# has_abstract=true, xml has no <back>.

def make_no_references():
    doc = new_doc()
    title = "Preliminary Notes on Blorf Topology and Snazzle Coherence"
    pp = doc.add_paragraph(style="Title")
    pp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = pp.add_run(title)
    r.font.name = "Times New Roman"; r.font.size = Pt(16); r.bold = True

    p(doc, "Quonkfield, D.", align=WD_ALIGN_PARAGRAPH.CENTER, italic=True)
    doc.add_paragraph()

    p(doc, "Abstract", style="Heading1", bold=True)
    p(doc,
      "These notes document preliminary observations on blorf topology as it "
      "relates to snazzle coherence in small-scale quonk networks. No formal "
      "methodology is claimed; findings are exploratory and intended to motivate "
      "future zibbit-aware research. No citation list is provided at this stage.")
    doc.add_paragraph()

    sections = [
        ("1. Observations", [
            "Initial blorf mapping reveals three dominant topology classes: "
            "linear, ring, and mesh. Snazzle coherence varies significantly "
            "across classes, with mesh exhibiting the highest stability.",
            "Quonk node degree appears to correlate with frumple throughput "
            "(r = 0.62) in preliminary measurements. Formal validation is pending.",
        ]),
        ("2. Next Steps", [
            "A controlled experiment with zibbit-instrumented nodes is planned "
            "for the next research cycle. Collaboration with the Blorfville "
            "Institute is under discussion.",
        ]),
    ]
    for heading, paras in sections:
        p(doc, heading, style="Heading1", bold=True)
        for body in paras:
            p(doc, body)
        doc.add_paragraph()

    # Deliberately NO references / bibliography section

    save(doc, "article_no_references.docx")


# ── Article 4: Review paper with "Works Cited" ───────────────────────────────
# Parser expectations: "Works Cited" heading → ref_count=8, has_abstract=true,
# section_count >= 5.

def make_review_paper():
    doc = new_doc()
    title = (
        "A Systematic Review of Frumple Architectures "
        "in Modern Zibbit Systems: Trends, Gaps, and Future Directions"
    )
    pp = doc.add_paragraph(style="Title")
    pp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = pp.add_run(title)
    r.font.name = "Times New Roman"; r.font.size = Pt(16); r.bold = True

    p(doc, "Snazzleton, A., Frumpleston, B., & Zibbitova, C.",
      align=WD_ALIGN_PARAGRAPH.CENTER, italic=True)
    doc.add_paragraph()

    p(doc, "Abstract", style="Heading1", bold=True)
    p(doc,
      "This systematic review synthesises 47 primary studies published between "
      "2015 and 2024 on frumple architectures deployed within zibbit-enabled "
      "systems. Following PRISMA guidelines, we identify four dominant "
      "architectural families: monolithic blorf, micro-snazzle, event-driven "
      "quonk, and hybrid zibbit-frumple. Cross-study meta-analysis reveals a "
      "33% average throughput advantage for hybrid architectures under bursty "
      "blorf loads. Critical gaps in multi-tenant quonk scalability and formal "
      "frumple verification methods are identified as priorities for future work.")
    p(doc, "Keywords: systematic review, frumple architecture, zibbit, PRISMA, quonk",
      italic=True, size=11)
    doc.add_paragraph()

    sections = [
        ("1. Introduction", [
            "Frumple architectures have evolved rapidly alongside the proliferation "
            "of zibbit-enabled hardware. Early systems relied on monolithic blorf "
            "designs that prioritised simplicity over scalability. The shift toward "
            "distributed quonk topologies has driven demand for more flexible "
            "frumple models, yet no comprehensive review exists.",
            "This review addresses that gap by applying a systematic methodology "
            "to the corpus of peer-reviewed literature published between 2015 and 2024.",
        ]),
        ("2. Review Protocol", [
            "Search strings combining terms 'frumple', 'zibbit', 'quonk', and "
            "'architecture' were applied to four databases: Snazzle Digital Library, "
            "Blorf Index, IEEE Xplorf, and ACM Quonk Library. After title/abstract "
            "screening and full-text review, 47 studies met inclusion criteria.",
            "Quality was assessed using the Zibbit Research Quality Rubric (ZRQR). "
            "Data were extracted independently by two reviewers; disagreements "
            "resolved by a third reviewer.",
        ]),
        ("3. Architectural Families", [
            "Monolithic blorf architectures (n=14) offer predictable latency but "
            "show poor scalability beyond 1000 quonk nodes. Micro-snazzle designs "
            "(n=12) achieve horizontal scaling at the cost of increased frumple "
            "coordination overhead.",
            "Event-driven quonk architectures (n=11) excel under bursty workloads "
            "but require careful snazzle buffer tuning. Hybrid zibbit-frumple "
            "systems (n=10) combine event-driven responsiveness with blorf isolation.",
        ]),
        ("4. Meta-Analysis Results", [
            "Hybrid architectures yield a pooled 33% throughput advantage over "
            "monolithic blorf (95% CI: 28-38%, I² = 41%). The advantage narrows "
            "to 19% under uniform quonk load, suggesting hybrid designs are most "
            "beneficial in bursty or heterogeneous environments.",
            "No statistically significant difference in frumple failure rate was "
            "found across the four families (χ² = 3.2, df = 3, p = 0.36), "
            "indicating that reliability is not a differentiating factor.",
        ]),
        ("5. Gaps and Future Directions", [
            "Multi-tenant quonk scalability is underexplored: only 3 of 47 studies "
            "evaluate frumple architectures in shared-infrastructure settings. "
            "Formal verification of snazzle safety properties is absent from all "
            "reviewed works.",
            "Future work should develop standardised zibbit benchmarks to enable "
            "cross-study comparisons, and should evaluate frumple architectures "
            "in edge-quonk deployments where network connectivity is intermittent.",
        ]),
        ("6. Conclusion", [
            "This review provides the first systematic synthesis of frumple "
            "architecture research in zibbit systems. Hybrid designs offer the "
            "strongest throughput profile for bursty quonk workloads. Scalability "
            "and formal verification remain open research problems.",
        ]),
    ]
    for heading, paras in sections:
        p(doc, heading, style="Heading1", bold=True)
        for body in paras:
            p(doc, body)
        doc.add_paragraph()

    # "Works Cited" — tests the third alias in references.rs
    p(doc, "Works Cited", style="Heading1", bold=True)
    for ref in [
        "Blorfman, Q. (2015). Monolithic blorf in early quonk deployments. "
        "Journal of Frumple Engineering, 1(1), 1-15.",
        "Snazzleton, M., & Wumble, T. (2017). Micro-snazzle scaling: promises "
        "and pitfalls. IEEE Trans. Quonk Syst., 3(2), 44-58.",
        "Zibbitova, C., Frumpleston, B., & Iqbal, A. (2019). Event-driven quonk "
        "architectures: a taxonomy. ACM Comput. Surv., 51(4), Art. 77.",
        "Quonkfield, D. (2020). Hybrid zibbit-frumple systems: design patterns. "
        "in Proc. ICFE 2020, pp. 201-215.",
        "Wumble, R. (2021). Formal verification of snazzle safety properties. "
        "J. Nonsensical Comput., 10(3), 300-318.",
        "Frumpleston, B., & Blorfman, Q. (2022). Multi-tenant quonk scalability: "
        "challenges and approaches. Cloud Frumple Rev., 8(1), 5-22.",
        "Snazzleton, A. (2023). Edge-quonk deployments under intermittent "
        "connectivity. IEEE IoQ Letters, 5(2), 88-95.",
        "Zibbitova, C. (2024). PRISMA extensions for zibbit system reviews. "
        "Systematic Review Methods, 3(1), 12-28.",
    ]:
        p(doc, ref)

    save(doc, "article_review_paper.docx")


# ── main ──────────────────────────────────────────────────────────────────────

if __name__ == "__main__":
    print("Generating test DOCX fixtures...")
    make_full_research()
    make_bibliography_alias()
    make_no_references()
    make_review_paper()
    print("Done.")
