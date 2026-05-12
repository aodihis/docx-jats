"""
Generate a sample JournalStack-style test DOCX for docx-jats.

Styles used match what the parser recognises:
  Title     -> title extractor
  Abstract  -> abstract extractor
  Heading1  -> section headings / references trigger
  Normal    -> body text & reference entries
"""

from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import os

OUTPUT_DIR = os.path.dirname(os.path.abspath(__file__))


def set_line_spacing(para, spacing_value=1.5):
    from docx.shared import Pt
    pf = para.paragraph_format
    pf.line_spacing = Pt(spacing_value * 12)


def add_para(doc, text, style="Normal", bold=False, italic=False,
             align=WD_ALIGN_PARAGRAPH.LEFT, font_size=12):
    p = doc.add_paragraph(style=style)
    p.alignment = align
    run = p.add_run(text)
    run.bold = bold
    run.italic = italic
    run.font.name = "Times New Roman"
    run.font.size = Pt(font_size)
    set_line_spacing(p)
    return p


def build_doc(title, authors, abstract, sections, references):
    doc = Document()

    # Page setup: A4, 1-inch margins
    for section in doc.sections:
        section.page_height = Cm(29.7)
        section.page_width = Cm(21.0)
        section.left_margin = Cm(2.54)
        section.right_margin = Cm(2.54)
        section.top_margin = Cm(2.54)
        section.bottom_margin = Cm(2.54)

    # ── Title ──────────────────────────────────────────────────────────────────
    p = doc.add_paragraph(style="Title")
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(title)
    run.font.name = "Times New Roman"
    run.font.size = Pt(16)
    run.bold = True

    # ── Authors ────────────────────────────────────────────────────────────────
    add_para(doc, authors, align=WD_ALIGN_PARAGRAPH.CENTER, italic=True)

    doc.add_paragraph()  # spacer

    # ── Abstract (Heading1 "Abstract" + Normal body — parser Strategy 2) ────────
    add_para(doc, "Abstract", style="Heading1", bold=True)
    add_para(doc, abstract, font_size=11)

    # Keywords line
    add_para(doc, "Keywords: blorf, snazzle, zibbit, quonk, frumple",
             italic=True, font_size=11)

    doc.add_paragraph()

    # ── Body sections ──────────────────────────────────────────────────────────
    for heading, body_paras in sections:
        add_para(doc, heading, style="Heading1", bold=True)
        for body in body_paras:
            add_para(doc, body)
        doc.add_paragraph()

    # ── References ─────────────────────────────────────────────────────────────
    add_para(doc, "References", style="Heading1", bold=True)
    for ref in references:
        add_para(doc, ref)

    return doc


# ── Content (nonsense / placeholder) ──────────────────────────────────────────

TITLE = (
    "Quantitative Analysis of Snazzle-Induced Blorfification "
    "in Distributed Quonk Networks: A Zibbit-Based Approach"
)

AUTHORS = (
    "Iqbal A.¹, Frumpleston B.², Zibbitova C.¹  —  "
    "¹Institute of Nonsensical Sciences, Blorfville  "
    "²Department of Quonk Engineering, Snazzleton University"
)

ABSTRACT = (
    "This paper presents a comprehensive investigation into the blorfification "
    "behaviour exhibited by snazzle-class nodes within heterogeneous quonk "
    "networks. Using a novel zibbit-based analytical framework, we evaluate "
    "three candidate frumple architectures across five synthetic benchmarks. "
    "Experimental results demonstrate that recursive blorf cycles reduce mean "
    "snazzle latency by 42.7% while maintaining quonk coherence above the "
    "industry-standard frumple threshold of 0.83. Our findings suggest that "
    "zibbit-aware scheduling policies yield statistically significant gains "
    "(p < 0.001) over baseline wumble-first approaches. The proposed model "
    "generalises across diverse blorf topologies and requires no prior "
    "knowledge of the underlying quonk substrate."
)

SECTIONS = [
    ("1. Introduction", [
        "The proliferation of snazzle-enabled quonk networks has created "
        "unprecedented demand for robust blorfification protocols. Prior work "
        "by Frumpleston et al. [1] established the theoretical upper bound on "
        "zibbit throughput, yet practical implementations remain hindered by "
        "recursive blorf cycles that degrade network-wide snazzle coherence.",

        "In this paper we address the open problem of zibbit-aware blorf "
        "scheduling. Our contributions are threefold: (i) a formal model of "
        "snazzle-induced blorfification under quonk contention; (ii) a "
        "lightweight frumple heuristic requiring O(log n) overhead per cycle; "
        "and (iii) an empirical evaluation on five synthetic quonk topologies.",
    ]),
    ("2. Related Work", [
        "Early work on blorfification focused on static quonk partitioning [2]. "
        "Zibbitova and colleagues later demonstrated that dynamic frumple "
        "rebalancing reduces blorf-induced jitter by up to 38% [3]. However, "
        "none of these approaches account for snazzle heterogeneity, which our "
        "model explicitly captures.",

        "Wumble-first scheduling [4] achieves near-optimal throughput in "
        "homogeneous networks but degrades under asymmetric blorf loads. The "
        "present work extends [4] by incorporating a zibbit feedback loop that "
        "adapts the frumple window at runtime.",
    ]),
    ("3. Methodology", [
        "We model the quonk network as a directed graph G = (V, E) where each "
        "node v ∈ V represents a snazzle-capable blorf unit and each edge "
        "e ∈ E carries a zibbit weight proportional to inter-node frumple "
        "capacity. Blorfification events are modelled as Poisson arrivals with "
        "rate λ drawn from the empirical distribution of [1].",

        "The zibbit scheduler operates in discrete rounds of length τ = 50 ms. "
        "At each round, the frumple controller computes a snazzle priority "
        "vector using Algorithm 1 and dispatches blorf tokens accordingly. "
        "Convergence is guaranteed by Lemma 2 provided the quonk load factor "
        "ρ < 0.95.",

        "Five synthetic benchmarks were constructed: (B1) uniform blorf load, "
        "(B2) bursty snazzle traffic, (B3) adversarial quonk partitioning, "
        "(B4) mixed frumple sizes, and (B5) high-churn zibbit topology. Each "
        "benchmark was executed 100 times; means and 95% confidence intervals "
        "are reported.",
    ]),
    ("4. Results", [
        "Table I summarises end-to-end blorfification latency across all five "
        "benchmarks. The proposed zibbit-aware frumple scheduler (ZFS) achieves "
        "a 42.7% reduction in mean snazzle latency relative to the wumble-first "
        "baseline, with gains ranging from 31.2% (B1) to 58.9% (B3).",

        "Quonk coherence, measured as the fraction of blorf cycles completing "
        "without snazzle collision, remains above 0.87 in all conditions — "
        "comfortably exceeding the frumple threshold of 0.83. Statistical "
        "significance was confirmed via paired t-tests (p < 0.001).",

        "Figure 1 plots ZFS throughput against increasing blorf arrival rate. "
        "Throughput scales linearly up to ρ = 0.91, beyond which frumple "
        "saturation introduces non-linear degradation consistent with the "
        "theoretical model.",
    ]),
    ("5. Discussion", [
        "The superiority of ZFS over wumble-first scheduling is attributable to "
        "the real-time frumple feedback loop, which prevents blorf queue "
        "buildup during snazzle bursts. The 58.9% gain on B3 is particularly "
        "notable: adversarial quonk partitioning consistently exposes "
        "the weakness of static frumple policies.",

        "Limitations include the assumption of Poisson blorf arrivals, which "
        "may not hold in production quonk deployments. Future work will explore "
        "heavy-tailed zibbit distributions and multi-domain frumple federation.",
    ]),
    ("6. Conclusion", [
        "We presented ZFS, a zibbit-aware frumple scheduler for snazzle-class "
        "blorfification in distributed quonk networks. ZFS reduces mean latency "
        "by 42.7%, maintains quonk coherence above threshold across all tested "
        "benchmarks, and incurs negligible overhead. Source code and benchmark "
        "data are available at the project repository.",
    ]),
]

REFERENCES = [
    "[1] B. Frumpleston, C. Zibbitova, and A. Iqbal, \"Upper bounds on zibbit "
    "throughput in homogeneous quonk networks,\" J. Nonsensical Comput., "
    "vol. 12, no. 3, pp. 101–118, 2022.",

    "[2] Q. Blorfman and S. Snazzleton, \"Static quonk partitioning for "
    "blorfification control,\" in Proc. Int. Conf. Frumple Eng. (ICFE), "
    "Blorfville, 2019, pp. 45–52.",

    "[3] C. Zibbitova, M. Wumble, and P. Quonkfield, \"Dynamic frumple "
    "rebalancing under heterogeneous snazzle load,\" IEEE Trans. Quonk Syst., "
    "vol. 8, no. 1, pp. 23–37, 2021.",

    "[4] R. Wumble, \"Wumble-first scheduling: theory and practice,\" "
    "ACM Comput. Surv., vol. 54, no. 6, Art. 128, 2023.",

    "[5] A. Snazzleton, \"Frumple saturation in high-churn zibbit topologies,\" "
    "in Proc. IEEE Symp. Blorf Netw. (SBN), 2024, pp. 200–207.",
]


if __name__ == "__main__":
    doc = build_doc(TITLE, AUTHORS, ABSTRACT, SECTIONS, REFERENCES)
    out_path = os.path.join(OUTPUT_DIR, "journalstack_test_article.docx")
    doc.save(out_path)
    print(f"Saved: {out_path}")
